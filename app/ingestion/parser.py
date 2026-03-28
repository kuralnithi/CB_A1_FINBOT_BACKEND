"""
Document parser — handles PDF, Markdown, DOCX, and CSV files.

Uses Docling for PDF/Markdown and python-docx for DOCX preprocessing.
CSV files are converted to Markdown tables.
"""
import os
import logging
import csv
from pathlib import Path
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


def convert_docx_to_markdown(filepath: str) -> str:
    """
    Convert a .docx file to Markdown text for Docling processing.
    Preserves headings, paragraphs, tables, and lists.
    """
    doc = DocxDocument(filepath)
    md_lines: list[str] = []

    for para in doc.paragraphs:
        style = para.style.name.lower() if para.style else ""
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue

        if "heading 1" in style:
            md_lines.append(f"# {text}")
        elif "heading 2" in style:
            md_lines.append(f"## {text}")
        elif "heading 3" in style:
            md_lines.append(f"### {text}")
        elif "heading 4" in style:
            md_lines.append(f"#### {text}")
        elif "list" in style:
            md_lines.append(f"- {text}")
        else:
            md_lines.append(text)
        md_lines.append("")

    # Process tables
    for table in doc.tables:
        md_lines.append("")
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        md_lines.append("| " + " | ".join(headers) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
        md_lines.append("")

    return "\n".join(md_lines)


def convert_csv_to_markdown(filepath: str) -> str:
    """Convert a CSV file to a Markdown table."""
    md_lines: list[str] = []

    with open(filepath, "r", encoding="utf-8-sig") as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        return ""

    # Header
    headers = rows[0]
    md_lines.append(f"# {Path(filepath).stem.replace('_', ' ').title()}")
    md_lines.append("")
    md_lines.append("| " + " | ".join(headers) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

    # Data rows
    for row in rows[1:]:
        # Pad row if needed
        padded = row + [""] * (len(headers) - len(row))
        md_lines.append("| " + " | ".join(padded[:len(headers)]) + " |")

    return "\n".join(md_lines)


def preprocess_file(filepath: str) -> tuple[str, str]:
    """
    Pre-process a file for Docling ingestion.

    For DOCX/CSV files, converts to a temporary Markdown file.
    For PDF/MD files, returns the original path.

    Returns:
        (processed_filepath, original_filename)
    """
    ext = Path(filepath).suffix.lower()
    original_filename = Path(filepath).name

    if ext == ".docx":
        logger.info(f"Converting DOCX to Markdown: {original_filename}")
        md_content = convert_docx_to_markdown(filepath)
        tmp_path = filepath.replace(".docx", "_converted.md")
        with open(tmp_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        return tmp_path, original_filename

    elif ext == ".csv":
        logger.info(f"Converting CSV to Markdown: {original_filename}")
        md_content = convert_csv_to_markdown(filepath)
        tmp_path = filepath.replace(".csv", "_converted.md")
        with open(tmp_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        return tmp_path, original_filename

    elif ext in (".pdf", ".md"):
        return filepath, original_filename

    else:
        logger.warning(f"Unsupported file type: {ext} for {filepath}")
        return "", original_filename


def scan_data_directory(data_dir: str) -> list[dict]:
    """
    Scan the data directory and return file info with collection mapping.

    Directory structure:
        data/
          general/    → collection: general,   access_roles: all roles
          finance/    → collection: finance,   access_roles: [finance, c_level]
          engineering/→ collection: engineering,access_roles: [engineering, c_level]
          marketing/  → collection: marketing, access_roles: [marketing, c_level]
          hr/         → collection: general,   access_roles: all roles
    """
    FOLDER_COLLECTION_MAP = {
        "general":     {"collection": "general",     "access_roles": ["employee", "finance", "engineering", "marketing", "c_level"]},
        "hr":          {"collection": "general",     "access_roles": ["employee", "finance", "engineering", "marketing", "c_level"]},
        "finance":     {"collection": "finance",     "access_roles": ["finance", "c_level"]},
        "engineering": {"collection": "engineering", "access_roles": ["engineering", "c_level"]},
        "marketing":   {"collection": "marketing",   "access_roles": ["marketing", "c_level"]},
    }

    files_info = []
    data_path = Path(data_dir)

    if not data_path.exists():
        logger.error(f"Data directory not found: {data_dir}")
        return files_info

    for subfolder in data_path.iterdir():
        if not subfolder.is_dir():
            continue

        folder_name = subfolder.name.lower()
        mapping = FOLDER_COLLECTION_MAP.get(folder_name)

        if not mapping:
            logger.warning(f"Unknown data subfolder '{folder_name}' — skipping")
            continue

        for file_path in subfolder.iterdir():
            if file_path.is_file() and not file_path.name.startswith("."):
                # Skip already-converted temp files
                if "_converted.md" in file_path.name:
                    continue

                files_info.append({
                    "filepath": str(file_path),
                    "filename": file_path.name,
                    "collection": mapping["collection"],
                    "access_roles": mapping["access_roles"],
                })

    logger.info(f"Found {len(files_info)} documents in {data_dir}")
    return files_info
